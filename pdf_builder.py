import re
import subprocess
import os
import platform
import shutil
import signal
import tempfile
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Load environment variables from .env file
load_dotenv()

def _app_base_dir() -> Path:
    return Path(__file__).resolve().parent


def get_soffice_path():
    """Get LibreOffice soffice command path for current OS."""
    system = platform.system()

    # Try common paths for each OS
    paths_to_try = []

    env_path = os.getenv("SOFFICE_PATH") or os.getenv("LIBREOFFICE_PATH")
    if env_path:
        paths_to_try.append(env_path)

    if system == "Darwin":  # macOS
        paths_to_try += [
            "/opt/homebrew/bin/soffice",
            "/usr/local/bin/soffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        ]
    elif system == "Linux":
        paths_to_try = [
            "/usr/bin/soffice",
            "/usr/local/bin/soffice",
        ]

    # Check if any path exists
    for path in paths_to_try:
        try:
            candidate = Path(path)
            if candidate.is_dir():
                executable = candidate / "soffice"
                if executable.exists():
                    return str(executable)
            elif candidate.exists():
                return str(candidate)
        except Exception:
            continue

    # Fall back to searching in PATH
    try:
        soffice = shutil.which("soffice")
        if soffice:
            return soffice
    except Exception:
        pass

    return None

# Configuration - use environment variables with fallback defaults.
# The style-reference template lives in assets/ (outside the resumes output
# dir) so it is portable and never confused with generated resumes or wiped
# when clearing output. Falls back to the legacy location for old setups.
_ASSET_TEMPLATE = _app_base_dir() / 'assets' / 'resume_template.docx'
_LEGACY_TEMPLATE = _app_base_dir() / 'resumes' / 'Tharun Manikonda Resume.docx'
DEFAULT_TEMPLATE = str(_ASSET_TEMPLATE if _ASSET_TEMPLATE.exists() else _LEGACY_TEMPLATE)
TEMPLATE_PATH = os.getenv("RESUME_TEMPLATE_PATH", DEFAULT_TEMPLATE)
BULLET = "●"
TEXT_W = 7.884   # usable width (A4 8.278" − 2 × 0.197")

FORMAT_PROFILES = {
    "outlook": {
        "font_name": None,
        "margins": None,
        "title_size": 14,
        "contact_size": 10,
        "body_size": 10,
        "summary_spacing_before": 4,
        "section_spacing_before": 4,
        "experience_gap": 5,
        "project_gap": 6,
        "bullet_spacing": 2,
        "text_width": TEXT_W,
    },
    "gmail": {
        "font_name": None,
        "margins": {
            "top": 0.25,
            "bottom": 0.25,
            "left": 0.28,
            "right": 0.28,
        },
        "title_size": 13,
        "contact_size": 9.5,
        "body_size": 9.8,
        "summary_spacing_before": 3,
        "section_spacing_before": 3,
        "experience_gap": 4,
        "project_gap": 5,
        "bullet_spacing": 1.4,
        "text_width": 7.718,
    },
}

# ── Helpers ────────────────────────────────────────────────────────────────────

def _format_profile(name) -> dict:
    return FORMAT_PROFILES.get((name or "outlook").lower(), FORMAT_PROFILES["outlook"])


def _set_font_name(r_pr, font_name: str) -> None:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), font_name)


def _apply_document_profile(doc, profile: dict) -> None:
    font_name = profile.get("font_name")
    if font_name:
        for style_name in ("Normal", "Body Text", "List Paragraph", "Heading 1", "Heading 2", "Title", "p1"):
            try:
                style = doc.styles[style_name]
            except KeyError:
                continue
            style.font.name = font_name
            _set_font_name(style._element.get_or_add_rPr(), font_name)

    # Left-align section headers (template defaults them centered). The name
    # and contact stay centered — set explicitly on those paragraphs.
    try:
        doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except (KeyError, AttributeError):
        pass

    margins = profile.get("margins")
    if margins:
        for section in doc.sections:
            section.top_margin = Inches(margins["top"])
            section.bottom_margin = Inches(margins["bottom"])
            section.left_margin = Inches(margins["left"])
            section.right_margin = Inches(margins["right"])


def _format_run(run, size=None, font_name=None):
    if size is not None:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
        _set_font_name(run._element.get_or_add_rPr(), font_name)
    return run

def _add_right_tab(para, pos_inches=TEXT_W):
    """Append a right-aligned tab stop to a paragraph."""
    pPr  = para._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(pos_inches * 1440)))
    tabs.append(tab)


def _runs(para, text, size=10, font_name=None):
    """
    Split text on **bold** markers and append correctly-formatted runs.
    Non-marked parts → normal weight; marked parts → bold.
    """
    for i, part in enumerate(re.split(r'\*\*(.*?)\*\*', text)):
        if part:
            r = para.add_run(part)
            _format_run(r, size=size, font_name=font_name)
            # Ensure bold is properly set for odd-numbered parts (marked text)
            if i % 2 == 1:  # Odd indices are **bold** text
                r.bold = True
                r.font.bold = True  # Double-ensure bold is applied
            else:
                r.bold = False
                r.font.bold = False


def _remove_table_borders(table) -> None:
    """Ensure a table renders with no visible borders (used for layout grids)."""
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')


def _tighten_cell_margins(cell, left_twips=0, right_twips=70) -> None:
    """Reduce a table cell's inner left/right padding so column text aligns
    with the rest of the document and the two columns sit close together."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn('w:tcMar'))
    if margins is None:
        margins = OxmlElement('w:tcMar')
        tc_pr.append(margins)
    for edge, val in (('left', left_twips), ('right', right_twips)):
        el = margins.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            margins.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')


def _spacer(doc, after_pt=3):
    """Insert a blank Body Text paragraph used as vertical spacer."""
    p = doc.add_paragraph(style='Body Text')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(after_pt)
    return p


def _set_hanging_indent(para, left_twips=86, hanging_twips=187):
    """Set hanging indent via OxmlElement using twips (for bullet wrapping).
    Left: 0.06" = 86 twips, Hanging: 0.13" = 187 twips
    """
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),    str(left_twips))
    ind.set(qn('w:hanging'), str(hanging_twips))
    pPr.append(ind)


def _set_compact_spacing(para):
    """Set spacing to zero and enable contextualSpacing (don't add space between same style)."""
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is not None:
        pPr.remove(sp)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'),  '0')
    pPr.append(sp)

    # Enable "Don't add space between paragraphs of same style"
    contextual = pPr.find(qn('w:contextualSpacing'))
    if contextual is not None:
        pPr.remove(contextual)
    contextual = OxmlElement('w:contextualSpacing')
    contextual.set(qn('w:val'), '1')
    pPr.append(contextual)


def _add_section_borders(para):
    """Add bottom border only to a section header paragraph (Heading 1)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    # Only bottom border
    el = OxmlElement('w:bottom')
    el.set(qn('w:val'),   'single')
    el.set(qn('w:sz'),    '6')         # 0.75 pt
    el.set(qn('w:space'), '1')
    el.set(qn('w:color'), '000000')
    pBdr.append(el)

    pPr.append(pBdr)


# ── Resume builder ─────────────────────────────────────────────────────────────

def build_resume_docx(resume_data: dict, output_docx: str, format_profile: str = "outlook") -> None:
    # Load reference template — inherits ALL styles, A4 page, 0.197" margins
    doc  = Document(TEMPLATE_PATH)
    profile = _format_profile(format_profile)
    font_name = profile.get("font_name")
    body_size = profile["body_size"]
    _apply_document_profile(doc, profile)
    body = doc.element.body

    # Clear every paragraph/table/SDT in the body; keep sectPr (page layout)
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    d = resume_data

    # ── NAME ──────────────────────────────────────────────────────────────
    p_name = doc.add_paragraph(d['name'], style='Title')
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(0)
    for run in p_name.runs:
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red color

    # ── CONTACT LINE ──────────────────────────────────────────────────────
    c  = d['contact']
    p_contact = doc.add_paragraph(style='Normal')
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after  = Pt(3)
    # Order: phone | email | github/location
    contact_str = f"{c['phone']} | {c['email']}"
    if c.get('github'):
        contact_str += f" | {c['github']}"
    elif c.get('location'):
        contact_str += f" | {c['location']}"
    r = p_contact.add_run(contact_str)
    _format_run(r, size=profile["contact_size"], font_name=font_name)

    # ── SUMMARY ───────────────────────────────────────────────────────────
    p_h = doc.add_paragraph('SUMMARY', style='Heading 1')
    p_h.paragraph_format.space_before = Pt(3)
    p_h.paragraph_format.space_after = Pt(3)
    _add_section_borders(p_h)
    for run in p_h.runs:
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red
        run.bold = True

    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    _runs(p, d['summary'], size=body_size, font_name=font_name)

    _spacer(doc, after_pt=3)

    # ── TECHNICAL SKILLS ──────────────────────────────────────────────────
    p_h = doc.add_paragraph('TECHNICAL SKILLS', style='Heading 1')
    p_h.paragraph_format.space_after = Pt(3)
    _add_section_borders(p_h)
    for run in p_h.runs:
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red
        run.bold = True

    # Render skill categories in two balanced columns to save vertical space.
    skills = [sk for sk in d['technical_skills'] if str(sk.get('category', '')).strip()]
    if skills:
        half = (len(skills) + 1) // 2  # left column takes the extra when odd
        columns = [skills[:half], skills[half:]]

        skills_table = doc.add_table(rows=1, cols=2)
        skills_table.autofit = False
        skills_table.allow_autofit = False
        gutter_in = 0.22
        col_w = (profile["text_width"] - gutter_in) / 2
        _remove_table_borders(skills_table)

        for col_idx, entries in enumerate(columns):
            cell = skills_table.cell(0, col_idx)
            cell.width = Inches(col_w)
            _tighten_cell_margins(cell)
            first = True
            for sk in entries:
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                p.style = doc.styles['p1']
                _set_compact_spacing(p)
                r1 = p.add_run(sk['category'])
                _format_run(r1, size=body_size, font_name=font_name)
                r1.bold = True
                r2 = p.add_run(f": {sk['items']}")
                _format_run(r2, size=body_size, font_name=font_name)

    # ── PROFESSIONAL EXPERIENCE ───────────────────────────────────────────
    p_h = doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Heading 1')
    p_h.paragraph_format.space_after = Pt(3)
    _add_section_borders(p_h)
    for run in p_h.runs:
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red
        run.bold = True

    for i, exp in enumerate(d['experience']):

        # Title (bold) [TAB] Dates on right
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(profile["experience_gap"] if i > 0 else 1.9)
        p.paragraph_format.space_after  = Pt(0)
        _add_right_tab(p, profile["text_width"])
        r1 = p.add_run(exp['title'])
        _format_run(r1, size=body_size, font_name=font_name)
        r1.bold = True
        p.add_run('\t')
        r2 = p.add_run(exp['dates'])
        _format_run(r2, size=body_size, font_name=font_name)

        # Company | Location
        p = doc.add_paragraph(f"{exp['company']} | {exp['location']}",
                               style='Heading 2')
        p.paragraph_format.space_before = Pt(1.2)
        p.paragraph_format.space_after = Pt(3)

        # Bullet points
        for b in exp['bullets']:
            p = doc.add_paragraph(style='List Paragraph')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(profile["bullet_spacing"])
            # Set hanging indent for bullet text wrapping
            _set_hanging_indent(p)  # Uses default: left=86 twips (0.06"), hanging=187 twips (0.13")
            r = p.add_run(f"{BULLET} ")
            _format_run(r, size=body_size, font_name=font_name)
            _runs(p, b, size=body_size, font_name=font_name)

    _spacer(doc, after_pt=3)

    # ── PROJECTS ──────────────────────────────────────────────────────────
    p_h = doc.add_paragraph('PROJECTS', style='Heading 1')
    p_h.paragraph_format.space_after = Pt(3)
    _add_section_borders(p_h)
    for run in p_h.runs:
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red
        run.bold = True

    for proj in d['projects']:
        p = doc.add_paragraph(proj['name'], style='Heading 2')
        p.paragraph_format.space_before = Pt(profile["project_gap"])
        p.paragraph_format.space_after = Pt(3)

        for b in proj['bullets']:
            p = doc.add_paragraph(style='Body Text')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(1.55)
            r = p.add_run(f"{BULLET} ")
            _format_run(r, size=body_size, font_name=font_name)
            _runs(p, b, size=body_size, font_name=font_name)

    _spacer(doc, after_pt=3)

    # ── EDUCATION ─────────────────────────────────────────────────────────
    p_h = doc.add_paragraph('EDUCATION', style='Heading 1')
    p_h.paragraph_format.space_after = Pt(3)
    _add_section_borders(p_h)
    for run in p_h.runs:
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red
        run.bold = True

    for edu in d['education']:
        p = doc.add_paragraph(edu['degree'], style='Heading 2')
        p.paragraph_format.space_before = Pt(1.9)
        p.paragraph_format.space_after = Pt(1)

        # Institution [TAB] Dates
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.053)
        _add_right_tab(p, profile["text_width"])
        r1 = p.add_run(edu['institution'])
        _format_run(r1, size=body_size, font_name=font_name)
        p.add_run('\t')
        r2 = p.add_run(edu['dates'])
        _format_run(r2, size=body_size, font_name=font_name)

    _spacer(doc, after_pt=3)

    # ── CERTIFICATIONS ────────────────────────────────────────────────────
    p_h = doc.add_paragraph('CERTIFICATIONS', style='Heading 1')
    p_h.paragraph_format.space_after = Pt(3)
    _add_section_borders(p_h)
    for run in p_h.runs:
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red
        run.bold = True

    for cert in d['certifications']:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(2.25)
        _set_hanging_indent(p)  # Uses default: left=86 twips (0.06"), hanging=187 twips (0.13")
        r1 = p.add_run(f"{BULLET} ")
        _format_run(r1, size=body_size, font_name=font_name)
        r2 = p.add_run(cert)
        _format_run(r2, size=body_size, font_name=font_name)

    doc.save(output_docx)
    print(f'  DOCX saved → {output_docx}')


def is_pdf_conversion_ready() -> tuple[bool, str]:
    """Check if PDF conversion tools are available (LibreOffice)."""
    soffice_path = get_soffice_path()

    if not soffice_path:
        return False, "LibreOffice not installed"

    proc = None
    try:
        proc = subprocess.Popen(
            [soffice_path, "--version"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            start_new_session=(os.name != "nt"),
        )
        stdout, stderr = proc.communicate(timeout=5)
        if proc.returncode == 0:
            return True, f"LibreOffice ready ({stdout.strip()})"
        return False, f"LibreOffice not responding: {stderr.strip()}"
    except subprocess.TimeoutExpired:
        if proc is not None:
            if os.name != "nt":
                os.killpg(proc.pid, signal.SIGKILL)
            else:
                proc.kill()
            proc.communicate()
        return False, "LibreOffice version check timed out"
    except Exception as e:
        return False, f"Error checking LibreOffice: {str(e)}"


def _convert_docx_to_pdf_via_libreoffice(docx_path: str, output_path: str, timeout_seconds: int = 120) -> None:
    """Convert DOCX to PDF using LibreOffice command-line."""
    try:
        print(f"  [PDF] Starting conversion: {docx_path}")

        # Get LibreOffice path for this OS
        soffice_path = get_soffice_path()
        print(f"  [PDF] Detected LibreOffice path: {soffice_path}")

        if not soffice_path:
            raise RuntimeError("LibreOffice not found")

        # Verify DOCX file exists
        if not os.path.exists(docx_path):
            raise RuntimeError(f"DOCX file not found: {docx_path}")
        print(f"  [PDF] DOCX file verified: {docx_path}")

        # Get output directory and filename
        output_dir = os.path.dirname(output_path)
        output_filename = os.path.basename(output_path)

        print(f"  [PDF] Output directory: {output_dir}")
        print(f"  [PDF] Output filename: {output_filename}")

        # Verify output directory exists
        if not os.path.exists(output_dir):
            raise RuntimeError(f"Output directory does not exist: {output_dir}")
        print(f"  [PDF] Output directory verified")

        # Use an isolated LibreOffice profile for each conversion. Shared
        # headless profiles can lock up under concurrent requests.
        user_installation_dir = Path(tempfile.mkdtemp(prefix="resume-tool-lo-"))
        cmd = [
            soffice_path,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--nolockcheck",
            f"-env:UserInstallation={user_installation_dir.as_uri()}",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ]

        print(f"  [PDF] Running command: {' '.join(cmd)}")
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            env=os.environ.copy(),
            start_new_session=(os.name != "nt"),
        )
        try:
            stdout, stderr = proc.communicate(timeout=timeout_seconds)
        except subprocess.TimeoutExpired as e:
            if os.name != "nt":
                os.killpg(proc.pid, signal.SIGKILL)
            else:
                proc.kill()
            stdout, stderr = proc.communicate()
            raise RuntimeError(
                f"LibreOffice conversion timed out after {timeout_seconds} seconds. "
                f"STDOUT: {stdout} STDERR: {stderr} Error: {str(e)}"
            )
        finally:
            shutil.rmtree(user_installation_dir, ignore_errors=True)

        print(f"  [PDF] Command return code: {proc.returncode}")
        if stdout:
            print(f"  [PDF] STDOUT: {stdout}")
        if stderr:
            print(f"  [PDF] STDERR: {stderr}")

        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice command failed with code {proc.returncode}. STDERR: {stderr}")

        # LibreOffice creates PDF with same name as DOCX but .pdf extension
        expected_pdf = os.path.join(output_dir, os.path.basename(docx_path).replace('.docx', '.pdf'))
        print(f"  [PDF] Checking for PDF at: {expected_pdf}")

        if not os.path.exists(expected_pdf):
            # List files in output directory for debugging
            files_in_dir = os.listdir(output_dir)
            raise RuntimeError(f"PDF not created at {expected_pdf}. Files in directory: {files_in_dir}")

        print(f"  [PDF] ✓ PDF created successfully: {expected_pdf}")

    except subprocess.TimeoutExpired as e:
        raise RuntimeError(f"LibreOffice conversion timed out after {timeout_seconds} seconds. Error: {str(e)}")
    except Exception as e:
        raise RuntimeError(f"LibreOffice PDF conversion failed: {str(e)}")



def convert_docx_to_pdf_via_libreoffice(docx_path: str, output_path: str, timeout_seconds: int = 180) -> None:
    print("  Converting to PDF via LibreOffice...")

    try:
        _convert_docx_to_pdf_via_libreoffice(docx_path, output_path, timeout_seconds)
        print(f'  PDF saved -> {output_path}')
    except Exception as e:
        # If conversion fails, raise error
        raise RuntimeError(f"PDF conversion failed: {str(e)}")


def build_resume_pdf(resume_data: dict, output_path: str, timeout_seconds: int = 180, format_profile: str = "outlook") -> None:
    """Build .docx then convert to PDF via LibreOffice."""
    docx_path = output_path.replace('.pdf', '.docx')
    build_resume_docx(resume_data, docx_path, format_profile=format_profile)
    convert_docx_to_pdf_via_libreoffice(docx_path, output_path, timeout_seconds=timeout_seconds)
